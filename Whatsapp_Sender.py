from lib2to3.pgen2 import driver
import edgedriver_autoinstaller
#from multiprocessing import process
#from multiprocessing.context import Process
from email import message
from logging import exception
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5 import QtWidgets, QtCore, QtGui
import sys
from PyQt5.uic import loadUi
import time
import multiprocessing
from multiprocessing.context import Process
from multiprocessing import process
from ctypes import c_char_p
from PyQt5.uic.uiparser import WidgetStack
#from pyrebase.pyrebase import Database
import pyrebase
import os
import socket
#source whatsapp imports below

import xlrd    
from selenium import webdriver  
from selenium.webdriver.support.ui import Select
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.by import By
from selenium.common.exceptions import TimeoutException
import pyautogui
import pyperclip
from io import BytesIO
#import ioS
import codecs
import win32clipboard
from PIL import Image
import docx




register_loaded = False
Logged_in = False
Register = None

DOCUMENT_Path = ""
Excel_path = ""
Photo_or_video_Path = ""


def update_label():
    global status
    global User_interface
    global status_of_internet
    status2 = str(status)
    status2 = status2.replace("Value(<class 'ctypes.c_char_p'>,", "").replace("'","").replace(")","")
    User_interface.value = widget.currentIndex()
    '''
    time.sleep(1)
    print(str(widget.currentIndex()))
    time.sleep(1)
    try:
        print(widget.indexOf("Register"))
    except Exception as I:
        print(I)
    '''

    if status_of_internet.value == "connected":
        try:
            if User_interface.value == 0:
                main.internet.setHidden(True)
            if User_interface.value == 1:
                Login.internet.setHidden(True)
            if User_interface.value == 2:
                Register.internet.setHidden(True)
                pass
        except Exception:
            pass
    elif status_of_internet.value == "not":
        try:
            if User_interface.value == 0:
                main.internet.setVisible(True)
            if User_interface.value == 1:
                Login.internet.setVisible(True)
            if User_interface.value == 2:
                Register.internet.setVisible(True)
        except Exception:
            pass




class Main(QMainWindow): #,FROM_MAIN):
    global DOCUMENT_Path
    global Excel_path
    global Photo_or_video_Path
    global Logged_in

    def __init__(self,parent=None):
        super(Main,self).__init__(parent)
        loadUi("Main_UI.ui",self)
        #self.setupUi(self)
        widget.setFixedSize(866,592)
        self.sign_out.clicked.connect(self.gotologin)
        #self.whatshesaid_2.setText("<font size=12 color='#4ac8eb'>"+ "listening" +"</font>")
        self.internet.setHidden(True)
        self.START.clicked.connect(self.start_sending)
        self.Message.clicked.connect(self.document_browse)
        self.Excel.clicked.connect(self.excel_browse)
        self.Attachment.clicked.connect(self.Photo_browse)
        self.START.clicked.connect(self.start_sending)

        
    def gotologin(self):
        global Logged_in
        if Logged_in == False:
            Login = login()
            widget.addWidget(Login)

        my_shine = open(r"C:\ProgramData\Whatsapp_sender_Logged_in.txt", "w")
        my_shine.write("N")
        my_shine.close()
        widget.setCurrentIndex(widget.currentIndex() +1)

    def document_browse(self):
        global DOCUMENT_Path
        fname = QFileDialog.getOpenFileName(self,"Open Docx", "Documents", "document files (*.docx)")
        DOCUMENT_Path = fname[0]
        #print(fname[0])

    def excel_browse(self):
        global Excel_path
        fname = QFileDialog.getOpenFileName(self,"Open xlsx", "Excel File With Phonenumbers", "excel files (*.xlsx)")
        Excel_path = fname[0]

    def Photo_browse(self):
        global Photo_or_video_Path
        fname = QFileDialog.getOpenFileName(self,"Open Photo/video", "Photo Or Video")
        Photo_or_video_Path = fname[0]


    def start_sending(self):
      global DOCUMENT_Path
      global Excel_path
      global Photo_or_video_Path
      def Read_docx(filepath):
          doc = docx.Document(filepath)
          FinalText = []
          for paragraph in doc.paragraphs:
              FinalText.append(paragraph.text)
        
          return '\n' .join(FinalText)

      try:
       i_Cant = False

       try:
           messege = str(Read_docx(DOCUMENT_Path))
           pyperclip.copy(messege)
       except:
           self.Error.setText("<font size=12 color='#ff0000'>"+"Message docx is empty or can't be read"+"</font>")
           self.Error.setHidden(False)
           i_Cant = True


       try:
           wait = self.wait_starting_chat.toPlainText()
           wait = int(wait)
       except:
           self.Error.setText("<font size=12 color='#ff0000'>"+"only type numbers in the waiting time field"+"</font>")
           self.Error.setHidden(False)
           i_Cant = True



       if DOCUMENT_Path == '':
           self.Error.setText("<font size=12 color='#ff0000'>"+"Please browse the docx containing message to send"+"</font>")
           self.Error.setHidden(False)

       elif Excel_path == '':
           self.Error.setText("<font size=12 color='#ff0000'>"+"Please browse the xlsx file with phonenumbers to send"+"</font>")
           self.Error.setHidden(False)

       elif i_Cant == True:
           print('i_cant')


       else:
           
        if Photo_or_video_Path == '':
            self.Error.setText("<font size=12 color='#ff0000'>"+"No photo/video browsed, only message will be sent"+"</font>")
            self.Error.setHidden(False)
            print('no photo')

        self.Error.setHidden(True)

        wait = self.wait_starting_chat.toPlainText()
        wait = int(wait)













        '''
        def Read_docx(filepath):
            doc = docx.Document(filepath)
            FinalText = []
            for paragraph in doc.paragraphs:
                FinalText.append(paragraph.text)
            
            return '\n' .join(FinalText)
        '''


        def send_to_clipboard(clip_type, data):
            win32clipboard.OpenClipboard()
            win32clipboard.EmptyClipboard()
            win32clipboard.SetClipboardData(clip_type, data)
            win32clipboard.CloseClipboard()

        #loc = (r'C:\Users\My Computer\Desktop\contacts.xls')   # THIS IS THE LOCATION OF THE XCEL FILE CONTAINING ALL THE CONTACTS
        loc = (Excel_path)

        ghs = 0  # THIS IS THE VARIABLE TO DEFINE THE COLUMB
        nsh = 0  # THIS VARIABLE IS USED A THE VALUE OF ROWS
        hf = 0   # THIS IS THE VARIABLE USED SO THAT THE WHILE LOOP CAN BE RAN MANY TIMES
        path = 1 # THIS IS THE VARIABLE ..... YOU WONT UNDERSTANT SO I AM NOT WRITING
        npf = 1  # THIS IS THE VARIABLE TO DEFINE HOW MANY TABS ARE OPENED RIGHT NOW + 1
        edgedriver_autoinstaller.install()
        browser = webdriver.Edge('msedgedriver.exe') #webdriver.Chrome('D:\\chromedriver.exe')  # THIS IS THE LOCATION OF YOUR WEBDRIVER
        wr=xlrd.open_workbook(loc) #opening excel file
        shee1 = wr.sheet_by_index(0) #sheet number
        num = shee1.cell_value(ghs,nsh) #cell value, ghs columb and row to get phone number
        num2 = str(num)  #TYPECASTING A FLOAT VALUE IN A STRING SO THAT IT CAN BE ADDED TO THE URL
        # Open function to open the file "MyFile1.txt" 
        # (same directory) in append mode and 


        #content = codecs.open(r"C:\Users\My Computer\Desktop\messege.txt", mode="r", encoding="utf-8")
        #content = open(r"C:\Users\My Computer\Desktop\messege.txt","r") 
        '''
        mesg = content.readlines()

        mesg = [sub.replace('\n', '') for sub in mesg]
        mesg = [sub.replace('\ufeff', '') for sub in mesg]
        mesg = [sub.replace('\r', '') for sub in mesg]
        '''

        #print(mesg)
        #messege = str(mesg)
        #messege = ""
        #messege = str(mesg)
        #print(messege)
        #messege = str(messege)
        messege = str(Read_docx(DOCUMENT_Path))
        pyperclip.copy(messege)
        browser.get('https://web.whatsapp.com/')
        #load_whatsapp = WebDriverWait(browser, 1000000000000).until(EC.presence_of_element_located((By.XPATH, '//*[@id="side"]/div[1]/div/label/div/div[2]')))
        load_whatsapp = WebDriverWait(browser, 1000000000000).until(EC.presence_of_element_located((By.XPATH, '//*[@id="app"]/div/div/div[3]/span/div/div')))
        #time.sleep(20)
        while hf <=  10000000:
            try:
                #messege = messege.replace("[","").replace("]","").replace("'","").replace("\ufeff","")
                #print(messege)
                pyperclip.copy(messege)
                num2 = shee1.cell_value(ghs,nsh)
                num2 = str(num2)
                num2 = num2[:-2]

                #patanahi = 'https://api.WhatsApp.com/send?phone=+91'+num2
                patanahi = 'https://wa.me/+91'+num2
                browser.execute_script("window.open('');")
                browser.switch_to.window(browser.window_handles[npf])
                
                browser.get(patanahi)
                
                print(patanahi)
                #browser.get('https://api.WhatsApp.com/send?phone=+91'+num2)
                coc = browser.find_element_by_xpath('//*[@id="action-button"]')
                load_coc = WebDriverWait(browser, 1000000).until(EC.presence_of_element_located((By.XPATH, '//*[@id="action-button"]'))) #maybe continue to chat
                coc = browser.find_element_by_xpath('//*[@id="action-button"]')
                coc.click()
                use_whatsapp_web = WebDriverWait(browser, 1000000).until(EC.presence_of_element_located((By.PARTIAL_LINK_TEXT, 'use WhatsApp Web')))
                continueto = browser.find_element_by_partial_link_text('use WhatsApp Web')
                continueto.click()
                #time.sleep(10)
                try:    
                        
                        myElem = WebDriverWait(browser, wait).until(EC.presence_of_element_located((By.XPATH, '//*[@id="main"]/footer/div[1]/div/span[2]/div/div[2]/div[1]/div/div[1]/p')))
                        typenum = browser.find_element_by_xpath('//*[@id="main"]/footer/div[1]/div/span[2]/div/div[2]/div[1]/div/div[1]/p')  # MESSEGE WALA TEXT BOX TO FIND KR RHE HAI  //*[@id="main"]/footer/div[1]/div/span[2]/div/div[2]/div[1]/div/div[2] old sending message
                        time.sleep(0.8)
                        use_here_buttom = browser.find_element_by_xpath('//*[@id="app"]/div/div/div/div/div/div/div[2]/div/div[2]')
                        use_here_buttom.click()
                        time.sleep(0.5)
                        typenum.click()

                        pyautogui.hotkey('ctrl', 'v')
                        #pyautogui.typewrite(messege)
                        typenum.send_keys(Keys.ENTER)
                        time.sleep(1)
                        #browser.find_elements_by_css_selector("[aria-label=Sent]")

                        pipicon = browser.find_element_by_xpath('//*[@id="main"]/footer/div[1]/div/span[2]/div/div[1]/div[2]/div/div/span')
                        pipicon.click()
                        time.sleep(1)
                        wait_for_photos_to_Open = WebDriverWait(browser, 10).until(EC.presence_of_element_located((By.XPATH, '//*[@id="main"]/footer/div[1]/div/span[2]/div/div[1]/div[2]/div/span/div[1]/div/ul/li[1]/button/span')))
                        phots = browser.find_element_by_xpath('//input[@accept="image/*,video/mp4,video/3gpp,video/quicktime"]')
                        #phots.click()
                        #filepath = r"C:\Users\My Computer\Desktop\Desert.mp4"
                        filepath = Photo_or_video_Path
                        phots.send_keys(filepath)
                        
                    

                        
                        #pyperclip.copy(r"C:\Users\My Computer\Desktop\Desert.jpg")
                        '''
                        image = Image.open(filepath)
                        output = BytesIO()
                        image.convert("RGB").save(output, "BMP")
                        data = output.getvalue()[14:]
                        output.close()

                        send_to_clipboard(win32clipboard.CF_DIB, data)                
                        
                        
                        '''
                        
                        #typenum.click()
                        #pyautogui.hotkey('ctrl', 'v')
                        #typenum.send_keys(Keys.ENTER)
                        to_wait_for_send_button = myElem = WebDriverWait(browser, 1000000).until(EC.presence_of_element_located((By.XPATH, '//*[@id="app"]/div[1]/div[1]/div[2]/div[2]/span/div[1]/span/div[1]/div/div[2]/div/div[2]/div[2]/div/div')))
                        sendbutton = browser.find_element_by_xpath('//*[@id="app"]/div[1]/div[1]/div[2]/div[2]/span/div[1]/span/div[1]/div/div[2]/div/div[2]/div[2]/div/div')
                        sendbutton.click()
                        time.sleep(15)

                        '''
                        if path == 1:
                            pyautogui.click(480,583)
                            pyautogui.typewrite('Desktop')
                            pyautogui.press('enter')
                            time.sleep(1)
                            pyautogui.click(480,583)
                            pyautogui.typewrite('Desert.jpg')
                            time.sleep(1)
                            pyautogui.press('enter')
                        else:
                            pyautogui.click(480,583)
                            pyautogui.typewrite('Desert.jpg')
                            time.sleep(1)
                            pyautogui.press('enter')
                        
                        time.sleep(2)
                        sendbutton = browser.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[2]/span/div/span/div/div/div[2]/span/div/div')
                        sendbutton.click()
                        time.sleep(4)
                        pyautogui.press('TAB')
                        pyautogui.press('TAB')
                        pyautogui.press('enter')
                        '''
                except Exception as identifier:
                    print(identifier)
            except Exception as identifier:
                print(identifier)
            #typenum = browser.find_element_by_xpath('//*[@id="main"]/footer/div[1]/div[2]/div/div[2]')
            #typenum.click()
            #typenum.send_keys('Bhopal Parisangh is inviting you to a scheduled Zoom meeting.\nTime: Sep 15, 2020 08:00 PM India\nJoin Zoom Meeting &  Go on this - https://bit.ly/2ZA1kiW\nMeeting ID: 847 5487 7095\nभोपाल परिसंघ की बैठक\nमंगलवार 15 सिंतबर 2020,  शाम 8 बजे')
            #typenum.send_keys(Keys.ENTER)
            #time.sleep(1)
            #pipicon = browser.find_element_by_xpath('//*[@id="main"]/footer/div[1]/div[1]/div[2]')
            #pipicon.click()
            #time.sleep(1)
            #phots = browser.find_element_by_xpath('//*[@id="main"]/footer/div[1]/div[1]/div[2]/span/div/div/ul/li[1]/button')
            #phots.click()
            #time.sleep(3)
            #if path == 1:
            #   pyautogui.click(480,583)
            #   pyautogui.typewrite('Desktop')
            #   pyautogui.press('enter')
            #   time.sleep(1)
            #   pyautogui.click(480,583)
            #   pyautogui.typewrite('Desert.jpg')
            #   time.sleep(1)
            #   pyautogui.press('enter')
            #else:
            #   pyautogui.click(480,583)
            #  pyautogui.typewrite('Desert.jpg')
            # time.sleep(1)
            #  pyautogui.press('enter')
            
            #time.sleep(2)
            #sendbutton = browser.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[2]/span/div/span/div/div/div[2]/span/div/div')
            #sendbutton.click()
            #time.sleep(1)
            npf = npf+1
            if npf == 3:
                browser.switch_to.window(browser.window_handles[0])
                time.sleep(0.5)
                browser.close()
                browser.switch_to.window(browser.window_handles[0])
                time.sleep(0.5)
                browser.close()
                npf = npf-2
                browser.switch_to.window(browser.window_handles[0])
                time.sleep(0.5)
                
            else:
                pass    

            ghs = ghs+1
            hf=hf+1
            path = path+1

      except Exception as A:
          print(A)

    


def Update_the_app(status_of_internet,User_interface):
    time.sleep(8)
    if status_of_internet.value=="connected":
        file = open(r"C:\ProgramData\WhatsApp_version.txt", 'r')
        current_version = file.read()
        file.close()
        print(current_version)
        database = Firebase.database()
        Server_Version = database.get()
        Server_version = str(Server_Version.val()).replace("OrderedDict", "").replace("(", "").replace(")" ,"").replace("[" ,"").replace("]" ,"").replace("'" ,"").replace("Beta" ,"").replace("," ,"").replace("Yup" ,"").replace(" " ,"").replace("Version","")
        print(Server_version)
        if current_version != Server_version:
            os.startfile('Whatsapp Sender updaterr.bat')
            os.system("TASKKILL /F /IM Whatsapp_Sender.exe")


        

def Internet(status_of_internet,User_interface):
    #os.startfile('AP.exe')
    time.sleep(3)
    #User_interface = widget.currentIndex()
    while True:
        time.sleep(1)
        try:
            socket.create_connection(("1.1.1.1", 53))
            status_of_internet.value = "connected"
        except Exception:
            status_of_internet.value = "not"




class login(QDialog):
    global register_loaded
    global Register
    def __init__(self):
        super(login,self).__init__()
        loadUi("Log_in.ui",self)
        widget.setFixedSize(866,592)
        self.Register_button.clicked.connect(self.switchtoregister)
        self.Login_button.clicked.connect(self.login_to_firebase)
        self.internet.setHidden(True)


    def switchtoregister(self):
        global register_loaded
        global Register
        if register_loaded == False:
            Register = register()
            widget.addWidget(Register)
        if register_loaded == True:
            pass
        #if widget.currentIndex == 
        widget.setCurrentIndex(widget.currentIndex() +1)
        #widget.removeWidget(self)
        #widget.setCurrentIndex(widget.currentIndex() +1)

    def login_to_firebase(self):
        global auth
        global Firebase
        email = self.Email_input.toPlainText()
        password = self.Password_input.toPlainText()
        try:
            login = auth.sign_in_with_email_and_password(email, password)
            #main = Main()
            #widget.addWidget(main)
            my_file = open(r"C:\ProgramData\Whatsapp_sender_Logged_in.txt", "w")
            email_credentials = email
            my_file.write(email_credentials +"\n")
            my_file.write(password)
            my_file.close()
            widget.setCurrentIndex(widget.currentIndex() - 1)
            widget.setFixedSize(866,592)
        except Exception as a:
            print(a)
            self.Hidden_label.setText("<font size=4 color='#ff0000'>"+"wrong or NOT registered email or WRONG password"+"</font>")



class register(QDialog):
    global register_loaded
    global auth
    global Firebase
    def __init__(self):
        super(register,self).__init__()
        loadUi("Register_user.ui",self)
        widget.setFixedSize(897,617)
        self.Login_Button.clicked.connect(self.gotologin)
        self.Register_button.clicked.connect(self.register_firebase)
        self.internet.setHidden(True)
    def gotologin(self):
        #Login = login()
        #widget.addWidget(Login)
        #if widget.currentIndex() == 4:
        #    widget.removeWidget(self)
        global register_loaded
        register_loaded = True
        widget.setCurrentIndex(widget.currentIndex() -1)

    def register_firebase(self):
        global auth
        global Firebase
        email = self.Email_user_input.toPlainText()
        password = self.Password_user_input.toPlainText()
        confirm_pass = self.Confirm_Passoword_user_input.toPlainText()
        if password == confirm_pass:
            try:
                user = auth.create_user_with_email_and_password(email,password)
                self.Hidden_thing.setText("<font size=5 color='#55ff00'>"+"you are registered, plaease login"+"</font>")
            except Exception as identifier:
                print(identifier)
                self.Hidden_thing.setText("<font size=5 color='#ff0000'>"+"wrong or registered email or weak password"+"</font>")
        else:
            self.Hidden_thing.setText("<font size=5 color='#ff0000'>"+"Passwords not matched"+"</font>")
            


firebaseConfig = {

    "apiKey": "AIzaSyBxCHXdgQAkbGygN5LAQuVCnDUFGSv3Dc8",
    "authDomain": "whatsapp-sender-7a064.firebaseapp.com",
    "projectId": "whatsapp-sender-7a064",
    "storageBucket": "whatsapp-sender-7a064.appspot.com",
    "messagingSenderId": "766117771099",
    "appId": "1:766117771099:web:22d45ad330e41d32ff4c7e",
    "measurementId": "G-BGM7ZM4CEX",
    "databaseURL": "https://whatsapp-sender-7a064-default-rtdb.firebaseio.com"
}

Firebase = pyrebase.initialize_app(firebaseConfig)
auth = Firebase.auth()
storage = Firebase.storage()
database = Firebase.database()

flags = QtCore.Qt.WindowFlags(QtCore.Qt.FramelessWindowHint)





if __name__ == '__main__':
    multiprocessing.freeze_support()
    app = QtWidgets.QApplication(sys.argv)
    try:
        file1 = open(r"C:\ProgramData\Whatsapp_sender_Logged_in.txt","r")
    except Exception:
        file1 = open(r"C:\ProgramData\Whatsapp_sender_Logged_in.txt","w+")
    login_state = file1.read()

    widget=QtWidgets.QStackedWidget()
    main = Main()
    widget.addWidget(main)
    #widget.setFixedSize(1381,801)
    widget.show()
    to_check = open(r"C:\ProgramData\Whatsapp_sender_Logged_in.txt")
    content = to_check.read()
    to_check.close()

    if content == "N":
        Login = login()
        widget.addWidget(Login)
        widget.setCurrentIndex(widget.currentIndex() +1)
    else:
        try:
            to_check3 = open(r"C:\ProgramData\Whatsapp_sender_Logged_in.txt")
            string_list = to_check3.readlines()
            to_check3.close()
            email_id = str(string_list[0])
            email_id = email_id.replace("\n","")
            print(email_id)
            #print(email_id)
            password_cre = str(string_list[1])
            logggggin = auth.sign_in_with_email_and_password(email_id,password_cre)
        except Exception as I:
            print(I)
            to_check2 = open(r"C:\ProgramData\Whatsapp_sender_Logged_in.txt", "w+")
            to_check2.write("N")
            to_check2.close()
            Login = login()
            widget.addWidget(Login)
            widget.setCurrentIndex(widget.currentIndex() +1)


    manager = multiprocessing.Manager()
    User_interface = manager.Value(c_char_p, 0)
    status_of_internet = manager.Value(c_char_p, "PRESS HOTKEY")


    p1 = Process(target=Internet, args=(status_of_internet, User_interface))
    p1.start()

    p2 = Process(target=Update_the_app, args=(status_of_internet, User_interface))
    p2.start()



    timer = QtCore.QTimer()
    timer.timeout.connect(update_label)

    sys.exit(app.exec_())
